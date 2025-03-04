import os
import logging
import asyncio
import threading
import openai
import telegram
import base64
from flask import Flask, request
from telegram.ext import Application, CommandHandler, MessageHandler, filters
from telegram.request import HTTPXRequest
from io import BytesIO

# Zusätzliche Bibliotheken für Dateiverarbeitung und -erstellung
try:
    import PyPDF2
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
except ImportError:
    PyPDF2 = None
try:
    import docx
except ImportError:
    docx = None
try:
    import pandas as pd
    import openpyxl
except ImportError:
    pd = None

# Umgebungsvariablen
TELEGRAM_BOT_TOKEN = os.getenv("TELEGRAM_BOT_TOKEN")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
WEBHOOK_URL = os.getenv("WEBHOOK_URL")  # z.B. "https://deinedomain.de/webhook"

# Logging konfigurieren
logging.basicConfig(
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s", level=logging.INFO
)
logger = logging.getLogger(__name__)

# Flask-App initialisieren
app = Flask(__name__)
openai.api_key = OPENAI_API_KEY

# Telegram Bot und Application initialisieren
request_instance = HTTPXRequest(pool_timeout=20)
bot = telegram.Bot(token=TELEGRAM_BOT_TOKEN, request=request_instance)
application = Application.builder().token(TELEGRAM_BOT_TOKEN).build()

# Globale Speicher
chat_histories = {}
doc_texts = {}

def get_chat_history(chat_id: str):
    if chat_id not in chat_histories:
        chat_histories[chat_id] = [{
            "role": "system",
            "content": ("You are an AI assistant for a Telegram bot. Answer concisely and helpfully. "
                        "Manchmal ironisch und frech.")
        }]
    return chat_histories[chat_id]

# OpenAI-Funktion: Generierung von Textantworten (GPT-4)
def generate_response(chat_id: str, message: str) -> str:
    history = get_chat_history(chat_id)
    history.append({"role": "user", "content": message})
    client = openai.OpenAI(api_key=OPENAI_API_KEY)
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=history,
        max_tokens=1500,
    )
    reply = response.choices[0].message.content.strip()
    history.append({"role": "assistant", "content": reply})
    return reply

# OpenAI-Funktion: Sprachgenerierung (Text-zu-Speech)
def generate_audio_response(text: str) -> bytes:
    client = openai.OpenAI(api_key=OPENAI_API_KEY)
    response = client.audio.speech.create(
        model="tts-1",
        voice="sage",
        input=text
    )
    return response.content

# OpenAI-Funktion: Sprachanalyse (Transkription via Whisper)
def transcribe_audio(audio_path: str) -> str:
    client = openai.OpenAI(api_key=OPENAI_API_KEY)
    with open(audio_path, "rb") as audio_file:
        transcription = client.audio.transcriptions.create(
            model="whisper-1",
            file=audio_file,
            response_format="text"
        )
    return transcription

# OpenAI-Funktion: Bildanalyse via Vision API
def analyze_image(image_path: str) -> str:
    client = openai.OpenAI(api_key=OPENAI_API_KEY)
    with open(image_path, "rb") as image_file:
        base64_image = base64.b64encode(image_file.read()).decode("utf-8")
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "user", "content": [
                {"type": "text", "text": "Was ist auf diesem Bild zu sehen?"},
                {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"}},
            ]},
        ],
        max_tokens=300,
    )
    return response.choices[0].message.content

# OpenAI-Funktion: Bilderstellung (DALL·E‑3)
def generate_image(prompt: str) -> str:
    client = openai.OpenAI(api_key=OPENAI_API_KEY)
    response = client.images.generate(
        model="dall-e-3",
        prompt=prompt,
        size="1024x1024",
        quality="standard",
        n=1,
    )
    return response.data[0].url

# Handler für den /start-Befehl
async def start(update, context):
    await update.message.reply_text(
        "Hallo! Ich bin dein AI-gestützter Telegram-Bot.\n"
        "Du kannst mir Nachrichten, Bilder, Sprachnachrichten oder Dokumente (PDF, DOCX, XLS/XLSX, TXT) senden.\n"
        "Mit /create <format> <Text-Befehl> kannst du Dateien (pdf, docx, xlsx, html) erstellen.\n"
        "Bei /create wird der Text-Befehl über die OpenAI API verarbeitet, um z.B. alle Monate zeilenweise auszugeben.\n"
        "Anschließend wird die Datei erstellt und an dich gesendet."
    )

# Handler für Textnachrichten
async def handle_message(update, context):
    chat_id = str(update.effective_chat.id)
    message = update.message.text
    if message.lower().startswith("erstelle ein bild von") or message.lower().startswith("generate an image of"):
        prompt = message.lower().replace("erstelle ein bild von", "").replace("generate an image of", "").strip()
        image_url = generate_image(prompt)
        get_chat_history(chat_id).append({"role": "user", "content": f"[Bildgenerierung] {prompt}"})
        get_chat_history(chat_id).append({"role": "assistant", "content": f"[Bild] {image_url}"})
        await context.bot.send_photo(chat_id=chat_id, photo=image_url)
    else:
        reply = generate_response(chat_id, message)
        await context.bot.send_message(chat_id=chat_id, text=reply)

# Handler für empfangene Fotos (Bildanalyse)
async def handle_photo(update, context):
    chat_id = str(update.effective_chat.id)
    photo = update.message.photo[-1]
    file = await bot.get_file(photo.file_id)
    image_path = f"temp_{photo.file_id}.jpg"
    await file.download_to_drive(image_path)
    
    description = analyze_image(image_path)
    os.remove(image_path)
    
    await context.bot.send_message(chat_id=chat_id, text=f"Bildanalyse: {description}")

# Handler für Sprachnachrichten (Voice-Input)
async def handle_voice(update, context):
    chat_id = str(update.effective_chat.id)
    voice = update.message.voice
    file = await bot.get_file(voice.file_id)
    audio_path = f"temp_{voice.file_id}.ogg"
    await file.download_to_drive(audio_path)
    
    text = transcribe_audio(audio_path)
    os.remove(audio_path)
    
    if "text" in text.lower():
        reply = generate_response(chat_id, text)
        await context.bot.send_message(chat_id=chat_id, text=reply)
    else:
        reply = generate_response(chat_id, text)
        audio_response = generate_audio_response(reply)
        with open("response.ogg", "wb") as audio_file:
            audio_file.write(audio_response)
        await context.bot.send_voice(chat_id=chat_id, voice=open("response.ogg", "rb"))
        os.remove("response.ogg")

# Handler für Dateiupload und -verarbeitung
async def handle_document(update, context):
    chat_id = str(update.effective_chat.id)
    document = update.message.document
    file_name = document.file_name
    file = await bot.get_file(document.file_id)
    local_path = f"temp_{document.file_id}_{file_name}"
    await file.download_to_drive(local_path)
    
    ext = file_name.split('.')[-1].lower()
    extracted_text = ""
    
    if ext == "pdf":
        if PyPDF2 is None:
            extracted_text = "PyPDF2 ist nicht installiert."
        else:
            try:
                with open(local_path, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    for page in reader.pages:
                        extracted_text += page.extract_text() + "\n"
            except Exception as e:
                extracted_text = f"Fehler beim Lesen der PDF: {e}"
    elif ext in ["doc", "docx"]:
        if docx is None:
            extracted_text = "python-docx ist nicht installiert."
        else:
            try:
                doc = docx.Document(local_path)
                for para in doc.paragraphs:
                    extracted_text += para.text + "\n"
            except Exception as e:
                extracted_text = f"Fehler beim Lesen des Dokuments: {e}"
    elif ext in ["xls", "xlsx"]:
        if pd is None:
            extracted_text = "pandas (und openpyxl) sind nicht installiert."
        else:
            try:
                df = pd.read_excel(local_path)
                extracted_text = df.to_csv(index=False)
            except Exception as e:
                extracted_text = f"Fehler beim Lesen der Excel-Datei: {e}"
    elif ext == "txt":
        try:
            with open(local_path, "r", encoding="utf-8") as f:
                extracted_text = f.read()
        except Exception as e:
            extracted_text = f"Fehler beim Lesen der Textdatei: {e}"
    else:
        extracted_text = "Dateiformat nicht unterstützt."
    
    doc_texts[chat_id] = extracted_text
    
    if extracted_text and "nicht unterstützt" not in extracted_text and not extracted_text.startswith("Fehler"):
        prompt = f"Fasse den folgenden Text zusammen:\n\n{extracted_text[:4000]}"
        summary = generate_response(chat_id, prompt)
    else:
        summary = extracted_text

    message = (f"Dokument '{file_name}' verarbeitet.\nZusammenfassung:\n{summary}\n\n"
               "Du kannst nun Fragen zum Dokument stellen mit /askdoc <deine Frage>.\n"
               "Falls du den vollständigen Text herunterladen möchtest, benutze /download_document.")
    await context.bot.send_message(chat_id=chat_id, text=message)
    
    os.remove(local_path)

# Handler für Fragen zum Dokument
async def handle_askdoc(update, context):
    chat_id = str(update.effective_chat.id)
    if chat_id not in doc_texts or not doc_texts[chat_id]:
        await context.bot.send_message(chat_id=chat_id, text="Es wurde noch kein Dokument verarbeitet.")
        return
    question = update.message.text.replace("/askdoc", "").strip()
    if not question:
        await context.bot.send_message(chat_id=chat_id, text="Bitte stelle deine Frage nach dem Befehl.")
        return
    prompt = (f"Beantworte folgende Frage basierend auf diesem Dokument:\n\n"
              f"Dokument:\n{doc_texts[chat_id][:4000]}\n\nFrage: {question}")
    answer = generate_response(chat_id, prompt)
    await context.bot.send_message(chat_id=chat_id, text=answer)

# Handler für Herunterladen des Dokuments
async def handle_download_document(update, context):
    chat_id = str(update.effective_chat.id)
    if chat_id not in doc_texts or not doc_texts[chat_id]:
        await context.bot.send_message(chat_id=chat_id, text="Es wurde noch kein Dokument verarbeitet.")
        return
    output_filename = f"processed_document_{chat_id}.txt"
    with open(output_filename, "w", encoding="utf-8") as out_file:
        out_file.write(doc_texts[chat_id])
    await context.bot.send_document(chat_id=chat_id, document=open(output_filename, "rb"), filename=output_filename)
    os.remove(output_filename)

# Handler für Bildgenerierung
async def handle_generate_image(update, context):
    chat_id = str(update.effective_chat.id)
    prompt = update.message.text.replace("/generate", "").strip()
    if not prompt:
        await context.bot.send_message(chat_id=chat_id, text="Bitte gib eine Bildbeschreibung an!")
        return
    image_url = generate_image(prompt)
    await context.bot.send_photo(chat_id=chat_id, photo=image_url)

# Neuer Handler: Dateierstellung basierend auf Texteingabe und OpenAI API
async def handle_create(update, context):
    chat_id = str(update.effective_chat.id)
    command_text = update.message.text.replace("/create", "").strip()
    
    if not command_text:
        await context.bot.send_message(
            chat_id=chat_id,
            text="Nutze: /create <format> <Text-Befehl>. Beispiel: '/create pdf Schreibe alle Monate in separate Zeilen'"
        )
        return
    
    parts = command_text.split(maxsplit=1)
    if len(parts) < 2:
        await context.bot.send_message(
            chat_id=chat_id,
            text="Bitte gib ein Format und einen Text-Befehl an, z. B. '/create pdf Schreibe alle Monate in separate Zeilen'"
        )
        return
    
    format_type, prompt = parts[0].lower(), parts[1]
    
    # Verarbeite den Textbefehl über die OpenAI API, um den Inhalt zu generieren
    file_content = generate_response(chat_id, prompt)
    
    temp_filename = f"output_{chat_id}.{format_type}"
    
    try:
        if format_type == "pdf":
            if not PyPDF2 or not hasattr(canvas, 'Canvas'):
                await context.bot.send_message(
                    chat_id=chat_id,
                    text="PDF-Erstellung nicht verfügbar (reportlab fehlt)."
                )
                return
            buffer = BytesIO()
            c = canvas.Canvas(buffer, pagesize=letter)
            # Schreibe den generierten Inhalt zeilenweise
            lines = file_content.split("\n")
            y = 750  # Start-Y-Position
            for line in lines:
                c.drawString(100, y, line)
                y -= 15  # Zeilenabstand
            c.showPage()
            c.save()
            buffer.seek(0)
            await context.bot.send_document(
                chat_id=chat_id,
                document=buffer,
                filename=temp_filename
            )
            buffer.close()
        elif format_type == "docx":
            if not docx:
                await context.bot.send_message(
                    chat_id=chat_id,
                    text="DOCX-Erstellung nicht verfügbar (python-docx fehlt)."
                )
                return
            document = docx.Document()
            document.add_paragraph(file_content)
            document.save(temp_filename)
            await context.bot.send_document(
                chat_id=chat_id,
                document=open(temp_filename, "rb"),
                filename=temp_filename
            )
            os.remove(temp_filename)
        elif format_type == "xlsx":
            if not pd or not openpyxl:
                await context.bot.send_message(
                    chat_id=chat_id,
                    text="XLSX-Erstellung nicht verfügbar (pandas/openpyxl fehlt)."
                )
                return
            df = pd.DataFrame({"Text": file_content.split("\n")})
            df.to_excel(temp_filename, index=False)
            await context.bot.send_document(
                chat_id=chat_id,
                document=open(temp_filename, "rb"),
                filename=temp_filename
            )
            os.remove(temp_filename)
        elif format_type == "html":
            html_content = f"""<!DOCTYPE html>
<html>
  <head>
    <meta charset="UTF-8">
    <title>Generated Document</title>
  </head>
  <body>
    <pre>{file_content}</pre>
  </body>
</html>"""
            with open(temp_filename, "w", encoding="utf-8") as f:
                f.write(html_content)
            await context.bot.send_document(
                chat_id=chat_id,
                document=open(temp_filename, "rb"),
                filename=temp_filename
            )
            os.remove(temp_filename)
        else:
            await context.bot.send_message(
                chat_id=chat_id,
                text="Ungültiges Format. Nutze: pdf, docx, xlsx, html"
            )
    except Exception as e:
        await context.bot.send_message(
            chat_id=chat_id,
            text=f"Fehler beim Erstellen der Datei: {str(e)}"
        )

# Globaler Event Loop
global_loop = asyncio.new_event_loop()
def start_loop(loop: asyncio.AbstractEventLoop):
    asyncio.set_event_loop(loop)
    loop.run_forever()
loop_thread = threading.Thread(target=start_loop, args=(global_loop,), daemon=True)
loop_thread.start()

# Handler registrieren
application.add_handler(CommandHandler("start", start))
application.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_message))
application.add_handler(MessageHandler(filters.PHOTO, handle_photo))
application.add_handler(MessageHandler(filters.VOICE, handle_voice))
application.add_handler(MessageHandler(filters.Document.ALL, handle_document))
application.add_handler(CommandHandler("generate", handle_generate_image))
application.add_handler(CommandHandler("askdoc", handle_askdoc))
application.add_handler(CommandHandler("download_document", handle_download_document))
application.add_handler(CommandHandler("create", handle_create))

# Webhook-Route
@app.route('/webhook', methods=['POST'])
def webhook():
    update_json = request.get_json(force=True)
    logger.info(f"Webhook erhalten: {update_json}")
    update = telegram.Update.de_json(update_json, bot)
    asyncio.run_coroutine_threadsafe(application.process_update(update), global_loop)
    return "OK", 200

# Home-Route
@app.route('/')
def home():
    return "Bot is running!", 200

if __name__ == '__main__':
    async def startup():
        await bot.initialize()
        await application.initialize()
        await bot.delete_webhook()
        success = await bot.set_webhook(WEBHOOK_URL)
        if success:
            logger.info(f"Webhook erfolgreich gesetzt: {WEBHOOK_URL}")
        else:
            logger.error("Webhook konnte nicht gesetzt werden!")
        await application.start()
    startup_future = asyncio.run_coroutine_threadsafe(startup(), global_loop)
    startup_future.result()
    app.run(host="0.0.0.0", port=int(os.environ.get("PORT", 5000)), debug=True)

